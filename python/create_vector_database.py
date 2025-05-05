import os
import sys
import glob
import docx
import PyPDF2
import re
import boto3
import io
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.docstore.document import Document
from langchain_openai import OpenAIEmbeddings
from langchain_community.vectorstores import FAISS
from dotenv import load_dotenv

# Load environment variables
load_dotenv()

def extract_text_from_docx(file_path):
    """Extract text from a .docx file"""
    try:
        doc = docx.Document(file_path)
        full_text = []
        for para in doc.paragraphs:
            full_text.append(para.text)
        return '\n'.join(full_text)
    except Exception as e:
        print(f"Error extracting text from {file_path}: {str(e)}")
        return ""

def extract_text_from_pdf(file_path):
    """Extract text from a .pdf file"""
    try:
        text = ""
        with open(file_path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            for page_num in range(len(pdf_reader.pages)):
                page = pdf_reader.pages[page_num]
                text += page.extract_text() + f"\n<<<PAGE_{page_num + 1}>>>\n"
        return text
    except Exception as e:
        print(f"Error extracting text from {file_path}: {str(e)}")
        return ""

# Initialize S3 client
def get_s3_client():
    """Get an S3 client using environment variables"""
    return boto3.client(
        's3',
        aws_access_key_id=os.environ.get('AWS_ACCESS_KEY_ID'),
        aws_secret_access_key=os.environ.get('AWS_SECRET_ACCESS_KEY'),
        region_name=os.environ.get('AWS_REGION', 'us-east-1')
    )

def extract_text_from_s3_docx(s3_client, bucket, key):
    """Extract text from a .docx file stored in S3"""
    try:
        response = s3_client.get_object(Bucket=bucket, Key=key)
        file_content = response['Body'].read()
        
        # Use BytesIO to create a file-like object
        docx_file = io.BytesIO(file_content)
        doc = docx.Document(docx_file)
        
        full_text = []
        for para in doc.paragraphs:
            full_text.append(para.text)
        return '\n'.join(full_text)
    except Exception as e:
        print(f"Error extracting text from S3 DOCX {key}: {str(e)}")
        return ""

def extract_text_from_s3_pdf(s3_client, bucket, key):
    """Extract text from a .pdf file stored in S3"""
    try:
        response = s3_client.get_object(Bucket=bucket, Key=key)
        file_content = response['Body'].read()
        
        # Use BytesIO to create a file-like object
        pdf_file = io.BytesIO(file_content)
        pdf_reader = PyPDF2.PdfReader(pdf_file)
        
        text = ""
        for page_num in range(len(pdf_reader.pages)):
            page = pdf_reader.pages[page_num]
            text += page.extract_text() + f"\n<<<PAGE_{page_num + 1}>>>\n"
        return text
    except Exception as e:
        print(f"Error extracting text from S3 PDF {key}: {str(e)}")
        return ""

def list_s3_files(s3_client, bucket, prefix="data/"):
    """List all files in an S3 bucket with the given prefix"""
    try:
        response = s3_client.list_objects_v2(Bucket=bucket, Prefix=prefix)
        if 'Contents' not in response:
            return []
        
        files = []
        for obj in response['Contents']:
            key = obj['Key']
            if key.lower().endswith(('.pdf', '.docx')):
                files.append({
                    'key': key,
                    'size': obj['Size'],
                    'last_modified': obj['LastModified']
                })
        return files
    except Exception as e:
        print(f"Error listing S3 files: {str(e)}")
        return []

def process_directory(directory_path, max_files=None):
    """Process all documents in a directory and its subdirectories, and in S3"""
    documents = []
    text_splitter = RecursiveCharacterTextSplitter(
        chunk_size=1000,
        chunk_overlap=200,
        length_function=len,
        separators=["\n\n", "\n", ".", " ", ""]
    )
    
    # Process local files
    docx_files = glob.glob(os.path.join(directory_path, "**/*.docx"), recursive=True)
    pdf_files = glob.glob(os.path.join(directory_path, "**/*.pdf"), recursive=True)
    
    local_files = docx_files + pdf_files
    if max_files:
        local_files = local_files[:max_files]
    
    print(f"Found {len(local_files)} local files to process")
    
    for i, file_path in enumerate(local_files):
        try:
            print(f"Processing local file {i+1}/{len(local_files)}: {file_path}")
            
            # Extract text based on file type
            if file_path.lower().endswith('.docx'):
                text = extract_text_from_docx(file_path)
            elif file_path.lower().endswith('.pdf'):
                text = extract_text_from_pdf(file_path)
            else:
                continue
            
            if not text:
                continue
            
            # Split text into chunks
            chunks = text_splitter.split_text(text)
            
            # Create documents
            for j, chunk in enumerate(chunks):
                doc = Document(
                    page_content=chunk,
                    metadata={
                        "source": file_path,
                        "file_name": os.path.basename(file_path),
                        "chunk_id": j,
                        "source_type": "local"
                    }
                )
                documents.append(doc)
                
        except Exception as e:
            print(f"Error processing local file {file_path}: {str(e)}")
    
    # Process S3 files
    try:
        s3_client = get_s3_client()
        bucket = os.environ.get('S3_BUCKET_NAME')
        
        if not bucket:
            print("S3_BUCKET_NAME environment variable not set, skipping S3 processing")
        else:
            s3_files = list_s3_files(s3_client, bucket)
            print(f"Found {len(s3_files)} S3 files to process")
            
            for i, file_info in enumerate(s3_files):
                try:
                    key = file_info['key']
                    print(f"Processing S3 file {i+1}/{len(s3_files)}: {key}")
                    
                    # Extract text based on file type
                    if key.lower().endswith('.docx'):
                        text = extract_text_from_s3_docx(s3_client, bucket, key)
                    elif key.lower().endswith('.pdf'):
                        text = extract_text_from_s3_pdf(s3_client, bucket, key)
                    else:
                        continue
                    
                    if not text:
                        continue
                    
                    # Split text into chunks
                    chunks = text_splitter.split_text(text)
                    
                    # Create documents
                    for j, chunk in enumerate(chunks):
                        doc = Document(
                            page_content=chunk,
                            metadata={
                                "source": key,
                                "file_name": os.path.basename(key),
                                "chunk_id": j,
                                "source_type": "s3",
                                "s3_key": key,
                                "s3_bucket": bucket
                            }
                        )
                        documents.append(doc)
                        
                except Exception as e:
                    print(f"Error processing S3 file {key}: {str(e)}")
    except Exception as e:
        print(f"Error setting up S3 client: {str(e)}")
    
    return documents

def create_vector_database(documents, output_dir, index_name="vector_database"):
    """Create or update a FAISS vector database from documents, merging with existing DB if present"""
    try:
        print(f"Creating/updating vector database with {len(documents)} new documents")
        embeddings = OpenAIEmbeddings()
        os.makedirs(output_dir, exist_ok=True)

        # Check if existing vector DB exists (both .pkl and FAISS index)
        pkl_path = os.path.join(output_dir, f"{index_name}.pkl")
        faiss_index_path = os.path.join(output_dir, f"{index_name}.faiss")
        vectorstore = None

        if os.path.exists(pkl_path) and os.path.exists(faiss_index_path):
            print("Existing vector database found. Loading and merging...")
            vectorstore = FAISS.load_local(
                folder_path=output_dir,
                embeddings=embeddings,
                index_name=index_name,
                allow_dangerous_deserialization=True
            )
            vectorstore.add_documents(documents)
        else:
            print("No existing vector database found. Creating new one.")
            vectorstore = FAISS.from_documents(documents, embeddings)

        # Save/overwrite both .pkl and FAISS index files
        vectorstore.save_local(output_dir, index_name)
        print(f"Vector database saved to {output_dir}/{index_name}.pkl and {output_dir}/{index_name}.faiss")
        return True
    except Exception as e:
        print(f"Error creating/updating vector database: {str(e)}")
        return False

def main():
    # Set paths
    data_dir = os.path.join(os.path.dirname(os.path.dirname(os.path.abspath(__file__))), "data")
    output_dir = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
    
    # Process documents
    print(f"Processing documents in {data_dir}")
    documents = process_directory(data_dir, max_files=100)  # Limit to 100 files for initial testing
    
    if not documents:
        print("No documents found or processed")
        return
    
    # Create vector database
    success = create_vector_database(documents, output_dir)
    
    if success:
        print("Vector database created successfully")
    else:
        print("Failed to create vector database")

if __name__ == "__main__":
    main()
